"""
Proposal Writer Tool for Agent Zero

Generates grant proposal sections using LLM and templates.
Ensures compliance with funder requirements.
"""

import os
from typing import Dict, List, Optional, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re


class ProposalWriter:
    """
    Proposal Writer Tool for Agent Zero
    
    Generates grant proposal sections tailored to funder requirements.
    Validates word/character counts and formatting rules.
    """
    
    def __init__(self, agent=None, **kwargs):
        """Initialize tool with agent context"""
        self.agent = agent
        self.name = "proposal_writer"
        self.args = kwargs
        
        # Section templates and prompts
        self.section_prompts = {
            'abstract': self._get_abstract_prompt(),
            'need': self._get_need_prompt(),
            'goals': self._get_goals_prompt(),
            'methods': self._get_methods_prompt(),
            'budget': self._get_budget_prompt(),
            'evaluation': self._get_evaluation_prompt(),
            'capacity': self._get_capacity_prompt()
        }
        
    async def execute(
        self,
        grant_id: str,
        section: str,
        org_profile: Dict[str, Any],
        action: str = "generate",
        existing_content: Optional[str] = None,
        requirements: Optional[Dict[str, Any]] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """
        Generate or refine proposal sections
        
        Args:
            grant_id: Grant opportunity ID
            section: Section name (abstract, need, goals, methods, budget, evaluation, capacity)
            org_profile: Organization information dict
            action: 'generate' or 'refine'
            existing_content: Current content (for refinement)
            requirements: Funder requirements (word limits, format, etc.)
            
        Returns:
            Dict with:
                - section_content: str (generated text)
                - word_count: int
                - character_count: int
                - compliance_status: Dict
                - suggestions: List[str]
        """
        
        try:
            # Validate section name
            if section not in self.section_prompts:
                return {
                    'error': f"Unknown section: {section}",
                    'valid_sections': list(self.section_prompts.keys())
                }
            
            # Get or set default requirements
            if not requirements:
                requirements = self._get_default_requirements(section)
            
            # Generate or refine content
            if action == "generate":
                content = await self._generate_section(
                    section=section,
                    org_profile=org_profile,
                    grant_id=grant_id,
                    requirements=requirements
                )
            elif action == "refine":
                content = await self._refine_section(
                    section=section,
                    existing_content=existing_content,
                    org_profile=org_profile,
                    requirements=requirements
                )
            else:
                return {'error': f"Unknown action: {action}. Use 'generate' or 'refine'"}
            
            # Calculate metrics
            word_count = len(content.split())
            char_count = len(content)
            
            # Check compliance
            compliance = self._check_compliance(
                content=content,
                requirements=requirements
            )
            
            # Generate suggestions
            suggestions = self._generate_suggestions(
                section=section,
                content=content,
                compliance=compliance,
                org_profile=org_profile
            )
            
            return {
                'section': section,
                'section_content': content,
                'word_count': word_count,
                'character_count': char_count,
                'compliance_status': compliance,
                'suggestions': suggestions,
                'message': f"Generated {section} section ({word_count} words)"
            }
            
        except Exception as e:
            return {
                'error': str(e),
                'message': f"Error generating {section}: {str(e)}"
            }
    
    async def _generate_section(
        self,
        section: str,
        org_profile: Dict[str, Any],
        grant_id: str,
        requirements: Dict[str, Any]
    ) -> str:
        """Generate content for a section"""
        
        # Get section-specific prompt
        base_prompt = self.section_prompts[section]
        
        # Build context from org profile
        org_context = self._build_org_context(org_profile)
        
        # Build full prompt
        full_prompt = f"""{base_prompt}

ORGANIZATION INFORMATION:
{org_context}

GRANT ID: {grant_id}

REQUIREMENTS:
- Maximum words: {requirements.get('max_words', 'no limit')}
- Required elements: {', '.join(requirements.get('required_elements', []))}

Generate a compelling, professional {section} section that addresses all requirements.
"""
        
        # If agent has LLM access, use it. Otherwise, generate template
        if self.agent and hasattr(self.agent, 'query_llm'):
            content = await self.agent.query_llm(full_prompt)
        else:
            # Generate template-based content (fallback)
            content = self._generate_template_content(section, org_profile)
        
        return content.strip()
    
    async def _refine_section(
        self,
        section: str,
        existing_content: str,
        org_profile: Dict[str, Any],
        requirements: Dict[str, Any]
    ) -> str:
        """Refine existing section content"""
        
        refine_prompt = f"""Review and improve this {section} section:

CURRENT CONTENT:
{existing_content}

REQUIREMENTS:
- Maximum words: {requirements.get('max_words', 'no limit')}
- Required elements: {', '.join(requirements.get('required_elements', []))}

IMPROVEMENTS NEEDED:
- Ensure all required elements are present
- Strengthen language and clarity
- Meet word count requirements
- Maintain professional tone

Provide refined version:
"""
        
        if self.agent and hasattr(self.agent, 'query_llm'):
            content = await self.agent.query_llm(refine_prompt)
        else:
            # Simple refinement: just return existing with minor cleanup
            content = self._clean_content(existing_content)
        
        return content.strip()
    
    def _generate_template_content(self, section: str, org_profile: Dict[str, Any]) -> str:
        """Generate template-based content (fallback when no LLM available)"""
        
        org_name = org_profile.get('name', '[Organization Name]')
        mission = org_profile.get('mission', '[Mission statement]')
        programs = org_profile.get('programs', [])
        
        templates = {
            'abstract': f"""{org_name} is a nonprofit organization dedicated to {mission}. This project will address critical community needs through innovative programming that serves [target population]. With a track record of success and strong community partnerships, {org_name} is positioned to effectively implement this initiative and achieve measurable outcomes. We request funding to expand our impact and reach more individuals in need.""",
            
            'need': f"""The community faces significant challenges that demand immediate attention. Research indicates [cite relevant statistics]. {org_name} has observed these challenges firsthand through our work with [describe population]. Current resources are insufficient to meet the growing demand for services. Without intervention, these challenges will continue to impact community wellbeing. This project addresses these needs by providing targeted support to those most affected.""",
            
            'goals': f"""Project Goal: [Primary goal statement]

Objectives:
1. Increase access to [services/resources] for [target population]
2. Improve [measurable outcome] by [percentage] over [timeframe]
3. Build capacity of [organization/community] to sustain program impact
4. Establish partnerships with [key stakeholders] to expand reach

These objectives align with {org_name}'s mission to {mission} and directly address the identified community needs.""",
            
            'methods': f"""{org_name} will implement this project through a phased approach:

Phase 1 - Planning and Preparation (Months 1-3):
- Assemble project team and establish governance structure
- Develop detailed implementation timeline
- Engage key stakeholders and partners

Phase 2 - Program Launch (Months 4-6):
- Begin service delivery to participants
- Establish data collection systems
- Monitor early indicators of success

Phase 3 - Full Implementation (Months 7-12):
- Scale program to full capacity
- Provide ongoing support to participants
- Conduct mid-course evaluation and adjustments

Our evidence-based approach draws on best practices from [relevant field] and has been successfully implemented in similar contexts.""",
            
            'budget': f"""Budget Summary:
Personnel: [Amount] - Project staff salaries and benefits
Program Supplies: [Amount] - Materials and resources for participants
Operations: [Amount] - Facilities, utilities, insurance
Evaluation: [Amount] - Data collection and analysis
Indirect Costs: [Amount] - Administrative support (15%)

Total Project Budget: [Total Amount]

{org_name} will provide matching funds of [amount] through [source], demonstrating organizational commitment to project success.""",
            
            'evaluation': f"""Evaluation Framework:
This project will use a mixed-methods evaluation approach combining quantitative and qualitative data.

Outcome Measures:
- [Metric 1]: Target - [percentage/number]
- [Metric 2]: Target - [percentage/number]
- [Metric 3]: Target - [percentage/number]

Data Collection Methods:
- Pre/post participant surveys
- Program attendance records
- Participant interviews and focus groups
- Partner feedback

Timeline:
- Quarterly progress reports
- Mid-term evaluation at Month 6
- Final evaluation at project completion

{org_name} will use evaluation findings to inform continuous improvement and demonstrate impact to stakeholders.""",
            
            'capacity': f"""{org_name} has a proven track record of successfully managing grant-funded projects. Established in [year], we have served over [number] individuals through programs including {', '.join(programs[:3]) if programs else '[program names]'}.

Organizational Strengths:
- Experienced leadership team with [number] years combined experience
- Strong financial management systems and clean audits
- Established partnerships with [key organizations]
- Board of Directors with diverse expertise
- History of meeting project milestones and deliverables

{org_name}'s annual budget of $[amount] demonstrates financial stability. We maintain [accreditation/certification] and follow best practices in nonprofit management."""
        }
        
        return templates.get(section, f"[{section.title()} section to be developed]")
    
    def _build_org_context(self, org_profile: Dict[str, Any]) -> str:
        """Build formatted organization context"""
        
        context_parts = []
        
        if 'name' in org_profile:
            context_parts.append(f"Organization: {org_profile['name']}")
        
        if 'mission' in org_profile:
            context_parts.append(f"Mission: {org_profile['mission']}")
        
        if 'org_type' in org_profile:
            context_parts.append(f"Type: {org_profile['org_type']}")
        
        if 'annual_budget' in org_profile:
            context_parts.append(f"Annual Budget: ${org_profile['annual_budget']:,}")
        
        if 'programs' in org_profile:
            programs = ', '.join(org_profile['programs'])
            context_parts.append(f"Programs: {programs}")
        
        if 'service_area' in org_profile:
            context_parts.append(f"Service Area: {org_profile['service_area']}")
        
        return '\n'.join(context_parts)
    
    def _get_default_requirements(self, section: str) -> Dict[str, Any]:
        """Get default requirements for section"""
        
        defaults = {
            'abstract': {'max_words': 250, 'required_elements': ['summary', 'goals', 'impact']},
            'need': {'max_words': 1000, 'required_elements': ['problem statement', 'data/evidence', 'target population']},
            'goals': {'max_words': 500, 'required_elements': ['objectives', 'outcomes', 'alignment']},
            'methods': {'max_words': 1500, 'required_elements': ['activities', 'timeline', 'approach']},
            'budget': {'max_words': 500, 'required_elements': ['line items', 'justification']},
            'evaluation': {'max_words': 750, 'required_elements': ['metrics', 'methods', 'timeline']},
            'capacity': {'max_words': 500, 'required_elements': ['experience', 'qualifications', 'resources']}
        }
        
        return defaults.get(section, {'max_words': 1000, 'required_elements': []})
    
    def _check_compliance(self, content: str, requirements: Dict[str, Any]) -> Dict[str, Any]:
        """Check if content meets requirements"""
        
        word_count = len(content.split())
        max_words = requirements.get('max_words')
        required_elements = requirements.get('required_elements', [])
        
        compliance = {
            'word_count_ok': True,
            'required_elements_present': [],
            'missing_elements': [],
            'is_compliant': True
        }
        
        # Check word count
        if max_words and word_count > max_words:
            compliance['word_count_ok'] = False
            compliance['is_compliant'] = False
        
        # Check required elements (simple keyword check)
        content_lower = content.lower()
        for element in required_elements:
            if element.lower() in content_lower:
                compliance['required_elements_present'].append(element)
            else:
                compliance['missing_elements'].append(element)
                compliance['is_compliant'] = False
        
        return compliance
    
    def _generate_suggestions(
        self,
        section: str,
        content: str,
        compliance: Dict[str, Any],
        org_profile: Dict[str, Any]
    ) -> List[str]:
        """Generate improvement suggestions"""
        
        suggestions = []
        
        if not compliance['word_count_ok']:
            suggestions.append("Content exceeds maximum word count. Consider condensing.")
        
        if compliance['missing_elements']:
            missing = ', '.join(compliance['missing_elements'])
            suggestions.append(f"Add required elements: {missing}")
        
        # Section-specific suggestions
        if section == 'abstract' and len(content.split()) < 100:
            suggestions.append("Abstract seems short. Ensure it covers key project elements.")
        
        if section == 'budget' and '$' not in content:
            suggestions.append("Include specific dollar amounts in budget section.")
        
        if section == 'evaluation' and not any(word in content.lower() for word in ['measure', 'metric', 'data']):
            suggestions.append("Strengthen evaluation section with specific metrics.")
        
        return suggestions
    
    def _clean_content(self, content: str) -> str:
        """Clean and format content"""
        
        # Remove extra whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Ensure proper sentence spacing
        content = re.sub(r'\.([A-Z])', r'. \1', content)
        
        return content.strip()
    
    # Section-specific prompt templates
    def _get_abstract_prompt(self) -> str:
        return """Write an executive summary/abstract that:
- Concisely describes the project and its significance
- States clear goals and expected outcomes
- Highlights organizational capacity
- Demonstrates community impact
Keep it compelling yet professional."""
    
    def _get_need_prompt(self) -> str:
        return """Write a statement of need that:
- Documents the problem with credible data
- Describes the target population and their challenges
- Explains why this issue requires attention now
- Shows organizational understanding of the need
Use statistics and evidence to support claims."""
    
    def _get_goals_prompt(self) -> str:
        return """Write goals and objectives that are:
- SMART (Specific, Measurable, Achievable, Relevant, Time-bound)
- Clearly aligned with identified needs
- Realistic given project scope and timeline
- Connected to measurable outcomes
Include both process and outcome objectives."""
    
    def _get_methods_prompt(self) -> str:
        return """Write a methods section that:
- Describes specific activities and approaches
- Provides a clear implementation timeline
- Explains the logic model connecting activities to outcomes
- Demonstrates evidence-based practices
- Details roles and responsibilities
Be concrete and detailed about implementation."""
    
    def _get_budget_prompt(self) -> str:
        return """Write a budget narrative that:
- Justifies each major expense category
- Shows cost-effectiveness
- Demonstrates financial feasibility
- Notes matching funds or leveraged resources
- Aligns expenses with proposed activities
Be transparent and realistic."""
    
    def _get_evaluation_prompt(self) -> str:
        return """Write an evaluation plan that:
- Identifies key performance indicators
- Describes data collection methods and tools
- Includes both process and outcome measures
- Provides a timeline for evaluation activities
- Explains how findings will be used
Be specific about what will be measured and how."""
    
    def _get_capacity_prompt(self) -> str:
        return """Write an organizational capacity section that:
- Demonstrates relevant experience and expertise
- Highlights past successes with similar projects
- Shows strong governance and financial management
- Describes qualified staff and leadership
- Notes key partnerships and resources
Build confidence in organizational ability to deliver."""


def export_to_docx(sections: Dict[str, str], filename: str) -> str:
    """Export proposal sections to Word document"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Grant Proposal', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add sections
    section_order = ['abstract', 'need', 'goals', 'methods', 'budget', 'evaluation', 'capacity']
    section_titles = {
        'abstract': 'Executive Summary',
        'need': 'Statement of Need',
        'goals': 'Goals and Objectives',
        'methods': 'Methods and Activities',
        'budget': 'Budget Narrative',
        'evaluation': 'Evaluation Plan',
        'capacity': 'Organizational Capacity'
    }
    
    for section in section_order:
        if section in sections:
            # Add section heading
            doc.add_heading(section_titles[section], 1)
            
            # Add section content
            para = doc.add_paragraph(sections[section])
            para.paragraph_format.line_spacing = 1.5
            
            # Add spacing
            doc.add_paragraph()
    
    # Save document
    doc.save(filename)
    return filename


# Tool metadata for Agent Zero registration
tool_info = {
    'name': 'proposal_writer',
    'description': 'Generate or refine grant proposal sections with funder compliance',
    'parameters': {
        'grant_id': 'Grant opportunity ID',
        'section': 'Section name (abstract, need, goals, methods, budget, evaluation, capacity)',
        'org_profile': 'Organization information dictionary',
        'action': '"generate" or "refine"',
        'existing_content': 'Current content (for refinement)',
        'requirements': 'Funder requirements (word limits, format rules)'
    },
    'example': 'proposal_writer(grant_id="ED-2025-001", section="abstract", org_profile={...}, action="generate")'
}
